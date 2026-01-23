import streamlit as st
import re
import random
from docx import Document
from io import BytesIO
from datetime import datetime, timedelta

st.set_page_config(page_title="İmtahan Hazırlayıcı", page_icon="📝")

@st.cache_data

def parse_docx(file):
    doc = Document(file)
    question_blocks = []
    paragraphs = list(doc.paragraphs)
    i = 0

    option_pattern = re.compile(r"^\s*[A-Ea-e][\)\.\:\-\s]+(.*)")
    question_pattern = re.compile(r"^\s*(\d+)\s*[.)]\s*(.*)")

    def is_numbered_paragraph(para):
        return para._p.pPr is not None and para._p.pPr.numPr is not None

    while i < len(paragraphs):
        para = paragraphs[i]
        text = ''.join(run.text for run in para.runs).strip()
        if not text:
            i += 1
            continue

        q_match = question_pattern.match(text)
        if q_match or is_numbered_paragraph(para):
            question_text = q_match.group(2).strip() if q_match else text.strip()
            i += 1
            options = []

            while i < len(paragraphs):
                option_text = ''.join(run.text for run in paragraphs[i].runs).strip()
                if not option_text:
                    i += 1
                    continue
                if question_pattern.match(option_text):
                    break
                match = option_pattern.match(option_text)
                if match:
                    options.append(match.group(1).strip())
                    i += 1
                else:
                    if len(options) < 5:
                        options.append(option_text)
                        i += 1
                    else:
                        break

            if len(options) >= 2:
                question_blocks.append((question_text, options))
        else:
            i += 1

    return question_blocks

@st.cache_data

def parse_open_questions(file):
    doc = Document(file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    questions = []
    for p in paragraphs:
        p = re.sub(r"^\s*\d+\s*[.)]?\s*", "", p)
        if p:
            questions.append(p)
    return questions

def create_shuffled_docx_and_answers(questions):
    new_doc = Document()
    answer_key = []

    for idx, (question, options) in enumerate(questions, start=1):
        new_doc.add_paragraph(f"{idx}) {question}")
        correct_answer = options[0]
        shuffled_options = options[:]
        random.shuffle(shuffled_options)

        for j, option in enumerate(shuffled_options):
            letter = chr(ord('A') + j)
            new_doc.add_paragraph(f"{letter}) {option}")
            if option.strip() == correct_answer.strip():
                answer_key.append(f"{idx}) {letter}")

    return new_doc, answer_key

if "page" not in st.session_state:
    st.session_state.page = "home"

if st.session_state.page == "home":
    st.title("📝 Testləri Qarışdır və Biliklərini Yoxla!")
    st.markdown("Zəhmət olmasa bir rejim seçin:")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        if st.button("📝 Özünü imtahan et"):
            st.session_state.page = "exam"
            st.rerun()
    with col2:
        if st.button("🎲 Sualları Qarışdır"):
            st.session_state.page = "shuffle"
            st.rerun()
    with col3:
        if st.button("🎫 Bilet İmtahanı"):
            st.session_state.page = "ticket"
            st.rerun()
    with col4:
        if st.button("🧮 Bal Hesablaması"):
            st.session_state.page = "score_calc"
            st.rerun()
    with col5:
        if st.button("ℹ️ İstifadə Qaydaları"):
            st.session_state.page = "help"
            st.rerun()
    
else:
    st.sidebar.title("⚙️ Menyu")
    if st.sidebar.button("🏠 Ana Səhifə"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.session_state.page = "home"
        st.rerun()

    menu = st.sidebar.radio("🔁 Rejimi dəyiş:", ["📝 Özünü İmtahan Et", "🎲 Sualları Qarışdır", "🎫 Bilet İmtahanı", "🧮 Bal Hesablaması", "ℹ️ İstifadə Qaydaları"],
                            index=["exam", "shuffle", "ticket", "score_calc", "help"].index(st.session_state.page))
    st.session_state.page = {"📝 Özünü İmtahan Et": "exam", "🎲 Sualları Qarışdır": "shuffle", "🎫 Bilet İmtahanı": "ticket", "🧮 Bal Hesablaması": "score_calc", "ℹ️ İstifadə Qaydaları": "help"}[menu]


if st.session_state.page == "exam":
    st.title("📝 Özünü Sına: İmtahan Rejimi ")
    uploaded_file = st.file_uploader("📤 İmtahan üçün Word (.docx) faylını seçin", type="docx")
    mode = st.radio(
        "📌 Sual seçimi:", 
        ["🔹 50 təsadüfi sual", "🔸 Bütün suallar", "🔻 Aralıqdan sual seçimi"],
        index=0
    )

    if uploaded_file:
        questions = parse_docx(uploaded_file)
        if not questions:
            st.error("❗ Heç bir sual tapılmadı.")
        else:
            if "exam_started" not in st.session_state:
                st.session_state.exam_started = False
            if "exam_submitted" not in st.session_state:
                st.session_state.exam_submitted = False
            if "exam_start_time" not in st.session_state:
                st.session_state.exam_start_time = None
            if "use_timer" not in st.session_state:
                st.session_state.use_timer = False

            selected = []

            if mode == "🔹 50 təsadüfi sual":
                selected = random.sample(questions, min(50, len(questions)))
                st.session_state.use_timer = True

            elif mode == "🔸 Bütün suallar":
                selected = questions
                st.session_state.use_timer = False

            elif mode == "🔻 Aralıqdan sual seçimi":
                st.markdown(f"💡 Faylda toplam **{len(questions)}** sual tapıldı.")
                start_q = st.number_input("🔢 Başlanğıc sual nömrəsi", min_value=1, max_value=len(questions), value=1, key="start_q")
                end_q = st.number_input("🔢 Sonuncu sual nömrəsi", min_value=start_q, max_value=len(questions), value=min(len(questions), start_q + 49), key="end_q")
            
                interval_questions = questions[start_q - 1:end_q]
            
                # Yeni seçim əlavə olunur
                available_modes = ["🔢 Ardıcıl", "🎲 Təsadüfi"]
                if len(interval_questions) >= 50:
                    available_modes.append("🎯 50 təsadüfi sual")
            
                order_mode = st.radio("📑 Sualların sıralanması və sayı:", available_modes, horizontal=True)
            
                if st.button("🚀 İmtahana Başla"):
                    if order_mode == "🎯 50 təsadüfi sual":
                        selected_questions = random.sample(interval_questions, 50)
                        st.session_state.use_timer = True
                    elif order_mode == "🎲 Təsadüfi":
                        selected_questions = random.sample(interval_questions, len(interval_questions))
                        st.session_state.use_timer = False
                    else:
                        selected_questions = interval_questions
                        st.session_state.use_timer = False
            
                    shuffled_questions = []
                    for q_text, opts in selected_questions:
                        correct = opts[0]
                        shuffled = opts[:]
                        random.shuffle(shuffled)
                        shuffled_questions.append((q_text, shuffled, correct))
            
                    st.session_state.exam_questions = shuffled_questions
                    st.session_state.exam_answers = [None] * len(shuffled_questions)
                    st.session_state.exam_start_time = datetime.now()
                    st.session_state.exam_started = True
                    st.rerun()



            if mode != "🔻 Aralıqdan sual seçimi" and not st.session_state.exam_started:
                if st.button("🚀 İmtahana Başla"):
                    shuffled_questions = []
                    for q_text, opts in selected:
                        correct = opts[0]
                        shuffled = opts[:]
                        random.shuffle(shuffled)
                        shuffled_questions.append((q_text, shuffled, correct))

                    st.session_state.exam_questions = shuffled_questions
                    st.session_state.exam_answers = [None] * len(shuffled_questions)
                    st.session_state.exam_start_time = datetime.now()
                    st.session_state.exam_started = True
                    st.rerun()

            elif st.session_state.exam_started and not st.session_state.exam_submitted:
                if st.session_state.get("use_timer", False):
                    elapsed = datetime.now() - st.session_state.exam_start_time
                    remaining = timedelta(minutes=60) - elapsed
                    seconds_left = int(remaining.total_seconds())

                    if seconds_left <= 0:
                        st.warning("⏰ Vaxt bitdi! İmtahan tamamlandı.")
                        st.session_state.exam_submitted = True
                        st.rerun()
                    else:
                        mins, secs = divmod(seconds_left, 60)
                        st.info(f"⏳ Qalan vaxt: {mins} dəq {secs} san")
                else:
                    st.info("ℹ️ Bu rejimdə zaman məhdudiyyəti yoxdur.")

                with st.form("exam_form"):
                    for i, (qtext, options, _) in enumerate(st.session_state.exam_questions):
                        st.markdown(f"**{i+1}) {qtext}**")
                        st.session_state.exam_answers[i] = st.radio("", options, key=f"q_{i}", label_visibility="collapsed")
                    submitted = st.form_submit_button("📤 İmtahanı Bitir")
                    if submitted:
                        st.session_state.exam_submitted = True
                        st.rerun()

            elif st.session_state.exam_submitted:
                st.success("🎉 İmtahan tamamlandı!")
                correct_list = [correct for _, _, correct in st.session_state.exam_questions]
                score = sum(1 for a, b in zip(st.session_state.exam_answers, correct_list) if a == b)
                total = len(correct_list)
                percent = (score / total) * 100

                st.markdown(f"### ✅ Nəticə: {score} düzgün cavab / {total} sual")
                st.markdown(f"<p style='font-size:16px;'>📈 Doğruluq faizi: <strong>{percent:.2f}%</strong></p>", unsafe_allow_html=True)
                st.progress(score / total)

                with st.expander("📊 Detallı nəticələr"):
                    for i, (ua, ca, (qtext, _, _)) in enumerate(zip(st.session_state.exam_answers, correct_list, st.session_state.exam_questions)):
                        status = "✅ Düzgün" if ua == ca else "❌ Səhv"
                        st.markdown(f"**{i+1}) {qtext}**\n• Sənin cavabın: {ua}\n• Doğru cavab: {ca} → {status}")

                if st.button("🔁 Yenidən Başla"):
                    keys_to_clear = [k for k in st.session_state if k.startswith("q_") or k in [
                        "exam_questions", "exam_answers", "exam_started", "exam_submitted", "exam_start_time", "use_timer"]]
                    for key in keys_to_clear:
                        st.session_state.pop(key)
                    st.rerun()



elif st.session_state.page == "shuffle":
    st.title("🎲 Test Suallarını Qarışdır və Cavab Açarı Yarat")
    uploaded_file = st.file_uploader("📤 Word (.docx) sənədini seçin", type="docx")
    mode = st.radio("💡 Sualların sayı:", ["🔹 50 təsadüfi sual", "🔸 Bütün suallar"], index=0)

    if uploaded_file:
        questions = parse_docx(uploaded_file)
        if len(questions) < 5:
            st.error("❗ Faylda kifayət qədər uyğun sual tapılmadı.")
        else:
            selected = random.sample(questions, min(50, len(questions))) if "50" in mode else questions
            new_doc, answer_key = create_shuffled_docx_and_answers(selected)

            output_docx = BytesIO()
            new_doc.save(output_docx)
            output_docx.seek(0)

            output_answers = BytesIO()
            output_answers.write('\n'.join(answer_key).encode('utf-8'))
            output_answers.seek(0)

            st.success("✅ Qarışdırılmış sənədlər hazırdır!")
            st.download_button("📥 Qarışdırılmış Suallar (.docx)", output_docx, "qarisdirilmis_suallar.docx")
            st.download_button("📥 Cavab Açarı (.txt)", output_answers, "cavab_acari.txt")

elif st.session_state.page == "ticket":
    st.title("🎫 Bilet İmtahanı (Açıq suallar)")
    uploaded_file = st.file_uploader("📤 Bilet sualları üçün Word (.docx) faylı seçin", type="docx")

    if uploaded_file:
        questions = parse_open_questions(uploaded_file)
        if len(questions) < 5:
            st.error("❗ Kifayət qədər sual yoxdur (minimum 5 tələb olunur).")
        else:
            if "ticket_questions" not in st.session_state:
                st.session_state.ticket_questions = []
            if "ticket_started" not in st.session_state:
                st.session_state.ticket_started = False

            if not st.session_state.ticket_started:
                if st.button("🎟️ Bilet Çək"):
                    st.session_state.ticket_questions = random.sample(questions, 5)
                    st.session_state.ticket_started = True

            if st.session_state.ticket_started:
                st.success("✅ Hazır bilet sualları:")
                for i, q in enumerate(st.session_state.ticket_questions, 1):
                    st.markdown(f"<p style='font-size:16px;'><strong>{i})</strong> {q}</p>", unsafe_allow_html=True)

                st.markdown("---")
                if st.button("🔁 Yenidən Bilet Çək"):
                    st.session_state.ticket_questions = random.sample(questions,5)
                    
elif st.session_state.page == "score_calc":
    st.title("🧮 Bal Hesablaması Sistemi")

    st.markdown("### Fənnin növünü seçin:")
    subject_type = st.radio("Bu fənn hansı əsasladır?", ["📘 Məşğələ", "🧪 Laboratoriya"])

    # Ümumi girişlər (kollekviumlar)
    a = st.number_input("1-ci kollekvium balı (maks 10)", min_value=0, max_value=10, step=1)
    b = st.number_input("2-ci kollekvium balı (maks 10)", min_value=0, max_value=10, step=1)
    c = st.number_input("3-cü kollekvium balı (maks 10)", min_value=0, max_value=10, step=1)

    if subject_type == "📘 Məşğələ":
        d = st.number_input("1-ci sərbəst iş balı (maks 5)", min_value=0, max_value=5, step=1)
        e = st.number_input("2-ci sərbəst iş balı (maks 5)", min_value=0, max_value=5, step=1)
        mesqele_orta = st.number_input("Məşğələ orta balı (maks 10)", min_value=0.0, max_value=10.0, step=0.1)
        l = st.number_input("Fənn üzrə dərs saatı (tam ədəd)", min_value=1, step=1)
        m = st.number_input("Neçə dəfə dərsdən qalmısınız (qayıb sayı)", min_value=0, max_value=l, step=1)

        if st.button("🔢 Balı Hesabla"):
            h = ((a + b + c) / 3) * 1.8        # Kollekvium: 18 bal
            i = d + e                         # Sərbəst işlər: 10 bal
            p = mesqele_orta * 1.2            # Məşğələ: 10 üzərindən daxil olunur → 12 bala miqyaslanır
            n = m * 2                         # Qayıb x2
            o = (n / l) * 10                  # Davamiyyət itkisi (maks 10)
            q = 10 - o                        # Davamiyyət balı
            k = h + i + p + q                 # Ümumi bal
            q_limit = l // 4 // 2             # İcazə verilən maksimum qayıb

            st.markdown("---")
            if m > q_limit:
                st.error("🚫 Sizin qayıb sayınız çox olduğundan imtahana buraxılmırsınız!")
            else:
                st.success(f"✅ İmtahan öncəsi topladığınız ümumi bal: **{k:.2f}**")

    elif subject_type == "🧪 Laboratoriya":
        d = st.number_input("1-ci sərbəst iş balı (maks 5)", min_value=0, max_value=5, step=1)
        e = st.number_input("2-ci sərbəst iş balı (maks 5)", min_value=0, max_value=5, step=1)
        f = st.number_input("Laboratoriya işlərinin ümumi sayı (tam ədəd)", min_value=1, step=1)
        g = st.number_input("Təhvil verilən laboratoriya sayı", min_value=0, max_value=f, step=1)
        l = st.number_input("Fənn üzrə dərs saatı (tam ədəd)", min_value=1, step=1)
        m = st.number_input("Neçə dəfə dərsdən qalmısınız (qayıb sayı)", min_value=0, max_value=l, step=1)

        if st.button("🔢 Balı Hesabla"):
            h = ((a + b + c) / 3) * 1.8        # Kollekvium: 18 bal
            i = d + e                         # Sərbəst işlər: 10 bal
            j = (g / f) * 12                  # Laboratoriya işləri: 12 bal
            n = m * 2                         # Qayıb x2
            o = (n / l) * 10                  # Davamiyyət itkisi (maks 10)
            p = 10 - o                        # Davamiyyət balı
            k = h + i + j + p                 # Ümumi bal
            q_limit = l // 4 // 2             # İcazə verilən maksimum qayıb

            st.markdown("---")
            if m > q_limit:
                st.error("🚫 Sizin qayıb sayınız çox olduğundan imtahana buraxılmırsınız!")
            else:
                st.success(f"✅ İmtahan öncəsi topladığınız ümumi bal: **{k:.2f}**")

                                                                      
elif st.session_state.page == "help":
    st.title("ℹ️ İstifadə Qaydaları və Yardım")
    st.markdown("""
**Xoş gəlmisiniz!** Bu proqram vasitəsilə müxtəlif formatlarda imtahan suallarını sınaqdan keçirə və özünüzü yoxlaya bilərsiniz. Aşağıda əsas funksiyalar, dəstəklənən fayl formatı və necə istifadə ediləcəyi barədə ətraflı məlumat verilmişdir:

---

### 📄 Dəstəklənən fayl formatı:
Yalnız `.docx` formatında Word sənədləri istifadə olunmalıdır.

### 📝 Test suallarının formatı:
- Hər bir sual nömrələnmiş olmalıdır:  
  `1) Bu bir nümunə sualdır?`
- Variantlar A-dan E-yə qədər olmalıdır:

  A) Doğru cavab  
  B) Yanlış cavab  
  C) Yanlış cavab  
  D) Yanlış cavab  
  E) Yanlış cavab

- **Diqqət**: Doğru cavab həmişə **birinci** yazılmalıdır (`A)` altında).

### 🧪 Rejimlər haqqında:

#### 📝 Özünü İmtahan Et:
- 3 rejim mövcuddur:  
  - **50 təsadüfi sual** (60 dəqiqəlik taymer ilə)  
  - **Bütün suallar** (vaxt məhdudiyyəti yoxdur)  
  - **Aralıqdan seçilmiş suallar** (istədiyiniz aralıqdan seçim; ardıcıl, təsadüfi və ya 50 təsadüfi variantları mövcuddur)  
- İmtahan zamanı cavablar qeyd olunur və sonunda nəticə, düzgün cavablar göstərilir.  
- Vaxt bitdikdə imtahan avtomatik tamamlanır.

#### 🎲 Sualları Qarışdır:
- Word sənədindən suallar və variantlar alınır, doğru cavablar qarışdırılır.  
- Qarışdırılmış suallar `.docx` faylı, cavab açarı isə `.txt` faylı kimi yüklənə bilər.  
- Suallar 50 təsadüfi və ya bütün suallar olaraq seçilə bilər.

#### 🎫 Bilet İmtahanı (Açıq suallar):
- Açıq tipli suallardan 5 təsadüfi sual seçilir.  
- "Bilet Çək" düyməsi ilə yeni suallar seçmək mümkündür.

### ⏱️ Vaxt məhdudiyyəti:
- **50 sual** rejimində 60 dəqiqəlik vaxt məhdudiyyəti var.  
- Digər rejimlərdə vaxt məhdudiyyəti yoxdur.

### 📤 Nəticələr:
- İmtahan tamamlandıqda düzgün və səhv cavablar, nəticə faizi göstərilir.  
- "Yenidən Başla" düyməsi ilə imtahan təkrar edilə bilər.

### 🧮 Bal Hesablanması Haqqında:

- **Məşğələ fənni üçün bal hesablanması:**

  - Kollekviumların ortalaması 18 baldır (3 kollekviumun orta balı * 1.8).  
  - 2 sərbəst iş balı toplanır (maksimum 10 bal).  
  - Məşğələ balı 10 üzərindən daxil edilir və 1.2 ilə vurularaq 12 bal hesablanır.  
  - Davamiyyət balı maksimum 10 baldır. Qayıb sayı dərs saatının müəyyən faizini keçərsə, bal azaldılır.  
  - Ümumi bal: kollekvium + sərbəst işlər + məşğələ + davamiyyət.  
  - Qayıb limiti: (Dərs saatı // 4) // 2, aşılması imtahana buraxılmamağa səbəb olur.

- **Laboratoriya fənni üçün bal hesablanması:**

  - Kollekviumların ortalaması 18 baldır.  
  - 2 sərbəst iş balı toplanır (maksimum 10 bal).  
  - Laboratoriya işlərinin sayı və təhvil verilən işlərin nisbətinə görə maksimum 12 bal hesablanır.  
  - Davamiyyət balı və qayıb limiti məşğələ fənni ilə eynidir.  
  - Ümumi bal: kollekvium + sərbəst işlər + laboratoriya + davamiyyət.

- **Bal hesablamasında maksimum ballar:**

  | Bal növü            | Maksimum bal |
  |---------------------|--------------|
  | Kollekvium          | 18           |
  | Sərbəst işlər       | 10           |
  | Məşğələ/Laboratoriya| 12           |
  | Davamiyyət          | 10           |
  | **Ümumi maksimum**  | **50**       |

- **Vacib qeyd:**  
  - Balın düzgün hesablanması üçün bütün girişlər maksimum hədlərdə olmalıdır.  
  - Qayıb sayı icazə verilən limitdən çox olarsa, imtahana buraxılmırsınız.

### ⚠️ Əgər:
- Fayl yüklədikdən sonra sual tapılmırsa, faylın strukturunu yoxlayın.  
- Sual və variantlar qarışmırsa, variantların düzgün `A)` formatında yazıldığından əmin olun.  
- Zaman bitərsə, imtahan avtomatik tamamlanacaq və nəticələr göstəriləcək.  
- Problemlə qarşılaşdıqda brauzeri yeniləyin və ya faylı yenidən yükləyin.

### 💡 Əlaqə və Yardım:
- Əgər çətinlik yaşayırsınızsa, bizimlə əlaqə saxlayın:  
  - Gmail: ismayilabbasov3032@gmail.com  
  - Mailru: ismayilabbasov3032@mail.ru

---

Uğurlar və uğurlu nəticələr!
""")
